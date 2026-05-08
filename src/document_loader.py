# import os
# import pdfplumber
# import docx

# def load_pdf(file_path: str) -> str:
#     """Extract text from a PDF file."""
#     text = ""
#     with pdfplumber.open(file_path) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#     return text.strip()

# def load_docx(file_path: str) -> str:
#     """Extract text from a DOCX file."""
#     doc = docx.Document(file_path)
#     text = "\n".join([para.text for para in doc.paragraphs])
#     return text.strip()

# def load_txt(file_path: str) -> str:
#     """Extract text from a TXT file."""
#     with open(file_path, "r", encoding="utf-8") as f:
#         text = f.read()
#     return text.strip()

# def load_document(file_path: str) -> str:
#     """Detect file type and extract text accordingly."""
#     if not os.path.exists(file_path):
#         raise FileNotFoundError(f"File not found: {file_path}")

#     ext = os.path.splitext(file_path)[1].lower()

#     if ext == ".pdf":
#         return load_pdf(file_path)
#     elif ext == ".docx":
#         return load_docx(file_path)
#     elif ext == ".txt":
#         return load_txt(file_path)
#     else:
#         raise ValueError(f"Unsupported file type: {ext}")

import os
import pdfplumber
import docx
from typing import Optional

# Optional OCR support for images
try:
    from PIL import Image
    import pytesseract
    OCR_ENABLED = True
except Exception:
    Image = None  # type: ignore
    pytesseract = None  # type: ignore
    OCR_ENABLED = False

def load_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()

def load_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

def load_txt(file_path: str) -> str:
    """Extract text from a TXT file."""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return text.strip()

def load_document(file_path: str) -> str:
    """Detect file type and extract text accordingly."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext == ".txt":
        return load_txt(file_path)
    elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
        if not OCR_ENABLED:
            raise RuntimeError("OCR support not available. Install Pillow and pytesseract and ensure Tesseract is installed on the system.")
        # perform OCR on the image and return extracted text
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            return text.strip()
        except Exception as e:
            raise RuntimeError(f"Failed to OCR image: {e}")
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def load_document_with_meta(file_path: str) -> dict:
    """Return extracted text and metadata for the given file.

    Returns:
        {
            "text": str,
            "pages": int,
            "file_type": str
        }
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text = load_pdf(file_path)
        try:
            with pdfplumber.open(file_path) as pdf:
                pages = len(pdf.pages)
        except Exception:
            pages = 0
        return {"text": text, "pages": pages, "file_type": "pdf"}
    elif ext == ".docx":
        text = load_docx(file_path)
        return {"text": text, "pages": 0, "file_type": "docx"}
    elif ext == ".txt":
        text = load_txt(file_path)
        lines = len(text.splitlines()) if text else 0
        return {"text": text, "pages": lines, "file_type": "txt"}
    elif ext in [".png", ".jpg", ".jpeg", ".tiff", ".bmp"]:
        if not OCR_ENABLED:
            raise RuntimeError("OCR support not available. Install Pillow and pytesseract and ensure Tesseract is installed on the system.")
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)
            return {"text": text.strip(), "pages": 1, "file_type": "image"}
        except Exception as e:
            raise RuntimeError(f"Failed to OCR image: {e}")
    else:
        raise ValueError(f"Unsupported file type: {ext}")
