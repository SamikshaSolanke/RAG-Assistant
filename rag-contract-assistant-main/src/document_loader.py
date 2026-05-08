# import os
# import pdfplumber
# import docx

# def load_pdf(file_path: str) -> str:
#     """Extract text from a PDF file."""
#     text = ""
#     with pdfplumber.open(file_path) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#     return text.strip()

# def load_docx(file_path: str) -> str:
#     """Extract text from a DOCX file."""
#     doc = docx.Document(file_path)
#     text = "\n".join([para.text for para in doc.paragraphs])
#     return text.strip()

# def load_txt(file_path: str) -> str:
#     """Extract text from a TXT file."""
#     with open(file_path, "r", encoding="utf-8") as f:
#         text = f.read()
#     return text.strip()

# def load_document(file_path: str) -> str:
#     """Detect file type and extract text accordingly."""
#     if not os.path.exists(file_path):
#         raise FileNotFoundError(f"File not found: {file_path}")

#     ext = os.path.splitext(file_path)[1].lower()

#     if ext == ".pdf":
#         return load_pdf(file_path)
#     elif ext == ".docx":
#         return load_docx(file_path)
#     elif ext == ".txt":
#         return load_txt(file_path)
#     else:
#         raise ValueError(f"Unsupported file type: {ext}")

import os
import pdfplumber
import docx

def load_pdf(file_path: str) -> str:
    """Extract text from a PDF file."""
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()

def load_docx(file_path: str) -> str:
    """Extract text from a DOCX file."""
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text.strip()

def load_txt(file_path: str) -> str:
    """Extract text from a TXT file."""
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return text.strip()

def load_document(file_path: str) -> str:
    """Detect file type and extract text accordingly."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    elif ext == ".txt":
        return load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
